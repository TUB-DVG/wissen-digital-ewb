import importlib

from django.core.management.base import BaseCommand, CommandError
from django.db.models import ManyToManyField, ForeignKey
from django.conf import settings
from docx import Document


class Command(BaseCommand):
    """ """

    def __init__(self):
        """ """
        super().__init__()
        self.document = Document()

    def _checkIfInInstalledApps(self, type_of_data):
        """Check if user given argument `type_of_data` matches
        one of the installed apps. If it does, check if a `data_import`-module
        is present in that app.
        """
        installed_django_apps = settings.INSTALLED_APPS
        app_names = [app.split(".")[0] for app in installed_django_apps]
        if type_of_data in app_names:
            if importlib.util.find_spec(type_of_data + ".models") is not None:
                return importlib.import_module(type_of_data + ".models")
        raise CommandError(
            """specified type_of_data has no corresponding app or has no 
            data_import.py in the app."""
        )

    def _checkIfTranslationConfFile(self, typeOfData):
        """ """
        return importlib.import_module(typeOfData + ".translation")

    def handle(
        self,
        *args: tuple,
        **options: dict,
    ) -> None:
        """This method is called from the manage.py when the 'data_import'
        command is called together with the manage.py.

        Parameters:
        *args:  tuple
            Tuple of arguments, which get unpacked.
        **options:  dict
            Dictionary, contains the parsed argument, given at the CLI
            when calling `python manage.py data_import`

        Returns:
        None
        """
        # filePathToData = options["pathCSV"][0]

        type_of_data = options["type_of_data"][0]
        nameOfModel = options["nameOfModel"][0]

        dataImportModule = self._checkIfInInstalledApps(type_of_data)
        translationModule = self._checkIfTranslationConfFile(type_of_data)

        allinstancesOfModel = self._getAllInstancesOfModel(
            dataImportModule, nameOfModel
        )
        modelClass = self._getModelClass(dataImportModule, nameOfModel)

        translationFieldList = self._getTranslationFields(
            translationModule, modelClass
        )

        for modelObj in allinstancesOfModel:
            self._checkForTranslationAttributes(modelObj, translationFieldList)
            # self.document.add_page_break()

        self.document.save(nameOfModel + ".docx")

    def _getModelClass(self, dataImportModule, nameOfModel):
        return getattr(dataImportModule, nameOfModel)

    def _getAllInstancesOfModel(self, dataImportModule, nameOfModel):
        """ """
        return self._getModelClass(dataImportModule, nameOfModel).objects.all()

    def _getTranslationFields(self, translationModule, modelClassObj):
        """ """
        return translationModule.translator.get_options_for_model(
            modelClassObj
        ).fields

    def _checkForTranslationAttributes(self, modelObj, translationFieldList):
        """ """
        self.document.add_heading(modelObj.__str__(), level=1)
        for translationField in translationFieldList:
            englishAttrForField = getattr(modelObj, translationField + "_en")
            if englishAttrForField != "" and englishAttrForField is not None:
                self.document.add_paragraph(
                    getattr(modelObj, translationField + "_de"),
                    style="List Bullet",
                )

                self.document.add_paragraph(
                    englishAttrForField, style="List Bullet"
                )

        modelFields = modelObj._meta.get_fields()

        for modelField in modelFields:
            try:
                relatedObj = getattr(modelObj, modelField.name)
            except:
                continue
            if isinstance(modelField, ForeignKey):

                try:
                    self.document.add_paragraph(
                        getattr(relatedObj, modelField.name + "_de"),
                        style="List Bullet",
                    )
                    self.document.add_paragraph(
                        getattr(relatedObj, modelField.name + "_en"),
                        style="List Bullet",
                    )
                except:
                    pass
            elif isinstance(modelField, ManyToManyField):
                try:
                    germanManyToManyValues = getattr(
                        relatedObj, modelField.name + "_de"
                    ).all()
                    germanValues = ""
                    for germanValue in germanManyToManyValues:
                        germanValues += germanValue + ", "
                    englishManyToManyValues = getattr(
                        relatedObj, modelField.name + "_en"
                    ).all()
                    englishValues = ""
                    for englishValue in englishManyToManyValues:
                        englishValues += englishValue + ", "
                    self.document.add_paragraph(
                        germanValues, style="List Bullet"
                    )
                    self.document.add_paragraph(
                        englishValues, style="List Bullet"
                    )
                except:
                    pass

        # else:
        #     raise CommandError(
        #         "Invalid file format. Please provide a .csv or .xlsx file.")
        # pathStr, filename = os.path.split(filePathToData)
        # self.filename = filename
        # self.targetFolder = options["targetFolder"][0]

        # appDataImportObj = data_import_module.DataImportApp(filePathToData)
        # appDataImportObj.personalDataFlag = options["personalData"]

        # header, data = appDataImportObj.load()
        # appDataImportObj.importList(header, data)

    def add_arguments(
        self,
        parser,
    ) -> None:
        """This method parses the arguments, which where given when
        calling the data_import-command together with the manage.py.
        The Arguments are then given to the handle-method, and can
        be accessed as python-variables.

        Parameters:
        parser: django.parser
        Django parser object, which handles the parsing of the command
        and arguments.

        Returns:
        None
        """
        parser.add_argument("type_of_data", nargs="+", type=str)
        parser.add_argument("nameOfModel", nargs="+", type=str)

        # parser.add_argument("pathCSV", nargs="+", type=str)
        # parser.add_argument(
        #     "--personalData",
        #     action="store_true",
        #     help="save personal data in the database",
        # )
